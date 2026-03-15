"""
export_controller.py — Export project data into various document formats.

Supported formats:
  - PDF   (.pdf)  — via reportlab
  - TXT   (.txt)  — plain-text outline
  - DOCX  (.docx) — Microsoft Word via python-docx
  - CSV   (.csv)  — flat table of all entities
  - ReqIF (.reqif) — Requirements Interchange Format (XML)

Every public function accepts a Project entity and a filesystem path,
recursively collects the hierarchy via get_children(), and writes
the output file.

Rich-text handling
──────────────────
QTextEdit produces full HTML documents.  Description and Body fields
are stored with a `%%PROJECT_MEDIA%%` placeholder for image paths.
The DOCX and PDF exporters parse this HTML to preserve bold, italic,
underline, lists, and embedded images in the output document.
"""

import csv
import html as html_mod
import os
import re
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import List, Optional

from controllers.db_controllers import get_children
from database.models import Entity


# ═══════════════════════════════════════════════════════════════════
# HELPER: recursive entity collection
# ═══════════════════════════════════════════════════════════════════

def _collect_tree(parent_id: int, depth: int = 0) -> list:
    """Return a flat list of (entity, depth) tuples in tree order."""
    result = []
    children = get_children(parent_id)
    for child in children:
        result.append((child, depth))
        result.extend(_collect_tree(child.id, depth + 1))
    return result


def _strip_html(text: Optional[str]) -> str:
    """Convert HTML to plain text by stripping tags."""
    if not text:
        return ""
    # Resolve the storage placeholder before stripping.
    try:
        from views.rich_text_editor import _html_from_storage
        text = _html_from_storage(text)
    except ImportError:
        pass
    # Extract body content if full HTML document.
    body_match = re.search(r"<body[^>]*>(.*)</body>", text, re.DOTALL | re.IGNORECASE)
    if body_match:
        text = body_match.group(1)
    # Replace <br>, <p>, <li> with newlines for readability.
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</li>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = html_mod.unescape(text)
    return text.strip()


def _resolve_html(text: Optional[str]) -> str:
    """Resolve storage placeholders and extract body content from rich HTML."""
    if not text:
        return ""
    try:
        from views.rich_text_editor import _html_from_storage
        text = _html_from_storage(text)
    except ImportError:
        pass
    body_match = re.search(r"<body[^>]*>(.*)</body>", text, re.DOTALL | re.IGNORECASE)
    if body_match:
        text = body_match.group(1)
    return text.strip()


def _entity_label(entity) -> str:
    """Return a display label like 'System' or 'Requirement'."""
    return entity.entity_type.capitalize()


# ═══════════════════════════════════════════════════════════════════
# RICH-TEXT HTML PARSER
# ═══════════════════════════════════════════════════════════════════

class _RichContentParser(HTMLParser):
    """Parse QTextEdit HTML into a list of structured content elements.

    Produces a list of dicts, each representing a paragraph or image:
      {"type": "paragraph", "runs": [{"text": "...", "bold": bool, "italic": bool, "underline": bool}, ...]}
      {"type": "image", "src": "/abs/path/to/img.png", "width": 300}
      {"type": "list_item", "list_kind": "ul"|"ol", "runs": [...]}
    """

    def __init__(self):
        super().__init__()
        self._elements = []
        self._current_runs = []
        self._bold = False
        self._italic = False
        self._underline = False
        self._in_list = None  # "ul" or "ol" or None
        self._in_li = False
        self._span_stack = []  # stack of (bold, italic, underline) before span

    def _style_has(self, style: str, prop: str, value: str) -> bool:
        """Check if a CSS style string contains a property:value pair."""
        if not style:
            return False
        for part in style.split(";"):
            part = part.strip().lower()
            if ":" in part:
                p, v = part.split(":", 1)
                if p.strip() == prop and value in v.strip():
                    return True
        return False

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        style = attrs_dict.get("style", "")

        if tag == "p":
            self._current_runs = []
        elif tag == "span":
            # Push current state onto stack.
            self._span_stack.append((self._bold, self._italic, self._underline))
            if self._style_has(style, "font-weight", "700") or self._style_has(style, "font-weight", "bold"):
                self._bold = True
            if self._style_has(style, "font-style", "italic"):
                self._italic = True
            if self._style_has(style, "text-decoration", "underline"):
                self._underline = True
        elif tag == "b" or tag == "strong":
            self._span_stack.append((self._bold, self._italic, self._underline))
            self._bold = True
        elif tag == "i" or tag == "em":
            self._span_stack.append((self._bold, self._italic, self._underline))
            self._italic = True
        elif tag == "u":
            self._span_stack.append((self._bold, self._italic, self._underline))
            self._underline = True
        elif tag == "ul":
            self._in_list = "ul"
        elif tag == "ol":
            self._in_list = "ol"
        elif tag == "li":
            self._in_li = True
            self._current_runs = []
        elif tag == "img":
            src = attrs_dict.get("src", "")
            # Strip file:/// prefix to get the local path.
            if src.startswith("file:///"):
                src = src[len("file:///"):]
            elif src.startswith("file://"):
                src = src[len("file://"):]
            width = 0
            try:
                width = int(attrs_dict.get("width", 0))
            except ValueError:
                pass
            self._elements.append({"type": "image", "src": src, "width": width or 300})
        elif tag == "br":
            self._current_runs.append({
                "text": "\n", "bold": False, "italic": False, "underline": False,
            })

    def handle_endtag(self, tag):
        if tag == "p":
            if self._current_runs:
                self._elements.append({"type": "paragraph", "runs": self._current_runs})
                self._current_runs = []
        elif tag in ("span", "b", "strong", "i", "em", "u"):
            if self._span_stack:
                self._bold, self._italic, self._underline = self._span_stack.pop()
        elif tag == "li":
            if self._current_runs:
                self._elements.append({
                    "type": "list_item",
                    "list_kind": self._in_list or "ul",
                    "runs": self._current_runs,
                })
                self._current_runs = []
            self._in_li = False
        elif tag in ("ul", "ol"):
            self._in_list = None

    def handle_data(self, data):
        if not data:
            return
        self._current_runs.append({
            "text": data,
            "bold": self._bold,
            "italic": self._italic,
            "underline": self._underline,
        })

    def get_elements(self) -> list:
        # Flush any remaining runs.
        if self._current_runs:
            self._elements.append({"type": "paragraph", "runs": self._current_runs})
            self._current_runs = []
        return self._elements


def _parse_rich_html(html_text: Optional[str]) -> list:
    """Parse DB-stored rich HTML into structured content elements."""
    resolved = _resolve_html(html_text)
    if not resolved:
        return []
    parser = _RichContentParser()
    parser.feed(resolved)
    return parser.get_elements()


# ═══════════════════════════════════════════════════════════════════
# TXT EXPORT
# ═══════════════════════════════════════════════════════════════════

def export_txt(project, filepath: str) -> None:
    """Export the project hierarchy as an indented plain-text outline."""
    lines = []
    lines.append(f"{'=' * 60}")
    lines.append(f"  {project.name}")
    if project.description:
        lines.append(f"  {_strip_html(project.description)}")
    lines.append(f"{'=' * 60}")
    lines.append("")

    tree = _collect_tree(project.id)
    for entity, depth in tree:
        indent = "    " * depth
        type_tag = f"[{_entity_label(entity)}]"

        if entity.entity_type == "requirement":
            req_id = getattr(entity, "req_id", "") or ""
            id_part = f"{req_id} — " if req_id else ""
            lines.append(f"{indent}{type_tag} {id_part}{entity.name}")
            body = _strip_html(getattr(entity, "body", None))
            if body:
                for bline in body.splitlines():
                    lines.append(f"{indent}    {bline}")
        else:
            lines.append(f"{indent}{type_tag} {entity.name}")
            desc = _strip_html(entity.description)
            if desc:
                for dline in desc.splitlines():
                    lines.append(f"{indent}    {dline}")

        lines.append("")

    Path(filepath).write_text("\n".join(lines), encoding="utf-8")


# ═══════════════════════════════════════════════════════════════════
# CSV EXPORT
# ═══════════════════════════════════════════════════════════════════

def export_csv(project, filepath: str) -> None:
    """Export all entities as a flat CSV table."""
    tree = _collect_tree(project.id)

    with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Depth", "Type", "Name", "Req ID", "Status",
            "Description", "Body", "Priority", "Rationale",
        ])

        for entity, depth in tree:
            writer.writerow([
                depth,
                _entity_label(entity),
                entity.name,
                getattr(entity, "req_id", "") or "",
                entity.status if entity.entity_type == "requirement" else "",
                _strip_html(entity.description),
                _strip_html(getattr(entity, "body", None)),
                getattr(entity, "priority", "") or "",
                _strip_html(getattr(entity, "rationale", None)),
            ])


# ═══════════════════════════════════════════════════════════════════
# DOCX EXPORT  (Microsoft Word) — with rich text and images
# ═══════════════════════════════════════════════════════════════════

def _docx_add_rich_content(doc, html_text: Optional[str], left_indent_pt=0,
                           default_color=None, default_size_pt=10):
    """Parse rich HTML and add formatted paragraphs/images to a Word doc."""
    from docx.shared import Pt, Inches, RGBColor

    elements = _parse_rich_html(html_text)
    if not elements:
        return

    color = default_color or RGBColor(100, 100, 100)

    for elem in elements:
        if elem["type"] == "image":
            src = elem["src"]
            if os.path.isfile(src):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(left_indent_pt)
                run = p.add_run()
                width_inches = min(elem.get("width", 300) / 96.0, 5.5)
                run.add_picture(src, width=Inches(width_inches))

        elif elem["type"] in ("paragraph", "list_item"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(left_indent_pt)

            if elem["type"] == "list_item":
                # Prefix with bullet or number marker.
                bullet = "\u2022 " if elem.get("list_kind") == "ul" else "- "
                marker_run = p.add_run(bullet)
                marker_run.font.size = Pt(default_size_pt)
                marker_run.font.color.rgb = color
                p.paragraph_format.left_indent = Pt(left_indent_pt + 12)

            for run_data in elem["runs"]:
                text = run_data["text"]
                if not text:
                    continue
                run = p.add_run(text)
                run.font.size = Pt(default_size_pt)
                run.font.color.rgb = color
                run.bold = run_data.get("bold", False)
                run.italic = run_data.get("italic", False)
                run.underline = run_data.get("underline", False)


def export_docx(project, filepath: str) -> None:
    """Export the project hierarchy as a formatted Word document."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title ─────────────────────────────────────────────────────
    title = doc.add_heading(project.name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project.description:
        _docx_add_rich_content(
            doc, project.description,
            default_color=RGBColor(150, 150, 150),
            default_size_pt=11,
        )

    doc.add_paragraph("")  # spacer

    # ── Recursive content ─────────────────────────────────────────
    tree = _collect_tree(project.id)
    for entity, depth in tree:
        if entity.entity_type == "requirement":
            _docx_add_requirement(doc, entity, depth)
        else:
            heading_level = min(depth + 1, 9)  # Word supports heading 1-9
            doc.add_heading(f"{entity.name}", level=heading_level)

            if entity.description:
                _docx_add_rich_content(
                    doc, entity.description,
                    left_indent_pt=depth * 18,
                    default_color=RGBColor(170, 170, 170),
                    default_size_pt=10,
                )

    doc.save(filepath)


def _docx_add_requirement(doc, entity, depth: int):
    """Add a requirement entry to the Word document."""
    from docx.shared import Pt, RGBColor

    req_id = getattr(entity, "req_id", "") or ""
    label = f"{req_id} — {entity.name}" if req_id else entity.name

    # Requirement heading-style paragraph.
    p = doc.add_paragraph()
    run = p.add_run(f"[REQ] {label}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(175, 122, 197)
    p.paragraph_format.left_indent = Pt(depth * 18)

    # Body — rich text with images.
    body = getattr(entity, "body", None)
    if body:
        _docx_add_rich_content(
            doc, body,
            left_indent_pt=depth * 18 + 12,
            default_color=RGBColor(80, 80, 80),
            default_size_pt=10,
        )


# ═══════════════════════════════════════════════════════════════════
# PDF EXPORT — with rich text and images
# ═══════════════════════════════════════════════════════════════════

def _pdf_rich_flowables(html_text: Optional[str], style, left_indent=0):
    """Convert rich HTML into a list of reportlab flowables (Paragraphs + Images)."""
    from reportlab.platypus import Paragraph, Image as RLImage
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm

    elements = _parse_rich_html(html_text)
    if not elements:
        return []

    flowables = []
    for elem in elements:
        if elem["type"] == "image":
            src = elem["src"]
            if os.path.isfile(src):
                img_width = min(elem.get("width", 300), 450)
                try:
                    img = RLImage(src, width=img_width, height=None)
                    # Preserve aspect ratio.
                    iw, ih = img.imageWidth, img.imageHeight
                    if iw and ih:
                        img.drawHeight = img_width * ih / iw
                    img.hAlign = "LEFT"
                    flowables.append(img)
                except Exception:
                    pass  # Skip broken images silently.

        elif elem["type"] in ("paragraph", "list_item"):
            # Build reportlab-safe markup from runs.
            markup_parts = []
            if elem["type"] == "list_item":
                bullet = "\u2022 " if elem.get("list_kind") == "ul" else "- "
                markup_parts.append(bullet)

            for run_data in elem["runs"]:
                text = run_data["text"]
                if not text:
                    continue
                # Escape XML special chars for reportlab.
                text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                text = text.replace("\n", "<br/>")
                if run_data.get("bold"):
                    text = f"<b>{text}</b>"
                if run_data.get("italic"):
                    text = f"<i>{text}</i>"
                if run_data.get("underline"):
                    text = f"<u>{text}</u>"
                markup_parts.append(text)

            markup = "".join(markup_parts)
            if markup.strip():
                para_style = ParagraphStyle(
                    f"rich_{id(elem)}",
                    parent=style,
                    leftIndent=left_indent + (12 if elem["type"] == "list_item" else 0),
                )
                flowables.append(Paragraph(markup, para_style))

    return flowables


def export_pdf(project, filepath: str) -> None:
    """Export the project hierarchy as a formatted PDF document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        HRFlowable,
    )

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Custom styles for each depth level.
    title_style = ParagraphStyle(
        "ProjectTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=HexColor("#333333"),
        spaceAfter=12,
        alignment=1,  # centre
    )

    def _heading_style(depth):
        sizes = {0: 16, 1: 14, 2: 12, 3: 11}
        colours = {0: "#2980b9", 1: "#27ae60", 2: "#f39c12", 3: "#8e44ad"}
        return ParagraphStyle(
            f"Heading_{depth}",
            parent=styles["Heading1"],
            fontSize=sizes.get(depth, 11),
            textColor=HexColor(colours.get(depth, "#555555")),
            leftIndent=depth * 18,
            spaceAfter=4,
            spaceBefore=10,
        )

    desc_style = ParagraphStyle(
        "Desc",
        parent=styles["Normal"],
        fontSize=9,
        textColor=HexColor("#888888"),
        spaceAfter=6,
    )

    req_header_style = ParagraphStyle(
        "ReqHeader",
        parent=styles["Normal"],
        fontSize=10,
        textColor=HexColor("#af7ac5"),
        spaceBefore=6,
        spaceAfter=2,
    )

    body_style = ParagraphStyle(
        "ReqBody",
        parent=styles["Normal"],
        fontSize=9,
        textColor=HexColor("#555555"),
        spaceAfter=8,
    )

    story = []

    # ── Title ─────────────────────────────────────────────────────
    story.append(Paragraph(project.name, title_style))
    if project.description:
        story.extend(_pdf_rich_flowables(
            project.description, desc_style, left_indent=0,
        ))
    story.append(HRFlowable(
        width="100%", thickness=1, color=HexColor("#6c5ce7"), spaceAfter=12
    ))

    # ── Recursive content ─────────────────────────────────────────
    tree = _collect_tree(project.id)
    for entity, depth in tree:
        if entity.entity_type == "requirement":
            req_id = getattr(entity, "req_id", "") or ""
            label = f"{req_id} &mdash; {entity.name}" if req_id else entity.name
            rstyle = ParagraphStyle(
                f"ReqH_{id(entity)}",
                parent=req_header_style,
                leftIndent=depth * 18,
            )
            story.append(Paragraph(f"<b>[REQ]</b> {label}", rstyle))

            body = getattr(entity, "body", None)
            if body:
                story.extend(_pdf_rich_flowables(
                    body, body_style, left_indent=depth * 18 + 12,
                ))
        else:
            hs = _heading_style(depth)
            story.append(Paragraph(entity.name, hs))
            if entity.description:
                story.extend(_pdf_rich_flowables(
                    entity.description, desc_style, left_indent=depth * 18 + 6,
                ))

    if not story:
        story.append(Paragraph("(empty project)", styles["Normal"]))

    doc.build(story)


# ═══════════════════════════════════════════════════════════════════
# ReqIF EXPORT  (Requirements Interchange Format — XML)
# ═══════════════════════════════════════════════════════════════════

def export_reqif(project, filepath: str) -> None:
    """Export the project as a ReqIF 1.0 XML document.

    Produces a simplified but valid ReqIF structure with:
    - One SPEC-OBJECT per entity
    - A SPECIFICATION with a SPEC-HIERARCHY mirroring the tree
    - ATTRIBUTE-DEFINITION for Name, Type, Description, ReqID, Body, Status
    """
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    ns = "http://www.omg.org/spec/ReqIF/20110401/reqif.xsd"

    root = ET.Element("REQ-IF", xmlns=ns)
    header = ET.SubElement(root, "THE-HEADER")
    rh = ET.SubElement(header, "REQ-IF-HEADER", IDENTIFIER="header")
    ET.SubElement(rh, "COMMENT").text = f"Exported from ReqMan — {project.name}"
    ET.SubElement(rh, "CREATION-TIME").text = now
    ET.SubElement(rh, "TITLE").text = project.name

    core = ET.SubElement(root, "CORE-CONTENT")
    content = ET.SubElement(core, "REQ-IF-CONTENT")

    # ── Data types ────────────────────────────────────────────────
    datatypes = ET.SubElement(content, "DATATYPES")
    ET.SubElement(
        datatypes, "DATATYPE-DEFINITION-STRING",
        IDENTIFIER="DT-String", LONG_NAME="String", MAX_LENGTH="10000",
    )

    # ── Spec-object type with attribute definitions ───────────────
    spec_types = ET.SubElement(content, "SPEC-TYPES")
    obj_type = ET.SubElement(
        spec_types, "SPEC-OBJECT-TYPE",
        IDENTIFIER="SOT-Entity", LONG_NAME="Entity",
    )
    attr_defs = ET.SubElement(obj_type, "SPEC-ATTRIBUTES")
    for attr_id, attr_name in [
        ("AD-Name", "Name"),
        ("AD-Type", "Type"),
        ("AD-Description", "Description"),
        ("AD-ReqID", "Requirement ID"),
        ("AD-Body", "Body"),
        ("AD-Status", "Status"),
    ]:
        ad = ET.SubElement(
            attr_defs, "ATTRIBUTE-DEFINITION-STRING",
            IDENTIFIER=attr_id, LONG_NAME=attr_name,
        )
        dt_ref = ET.SubElement(ad, "TYPE")
        ET.SubElement(dt_ref, "DATATYPE-DEFINITION-STRING-REF").text = "DT-String"

    # Specification type.
    ET.SubElement(
        spec_types, "SPECIFICATION-TYPE",
        IDENTIFIER="ST-Spec", LONG_NAME="Specification",
    )

    # ── Spec objects ──────────────────────────────────────────────
    spec_objects = ET.SubElement(content, "SPEC-OBJECTS")
    tree = _collect_tree(project.id)
    for entity, depth in tree:
        so = ET.SubElement(
            spec_objects, "SPEC-OBJECT",
            IDENTIFIER=f"SO-{entity.id}", LONG_NAME=entity.name,
        )
        vals = ET.SubElement(so, "VALUES")
        _reqif_add_attr(vals, "AD-Name", entity.name)
        _reqif_add_attr(vals, "AD-Type", _entity_label(entity))
        _reqif_add_attr(vals, "AD-Description", _strip_html(entity.description))
        if entity.entity_type == "requirement":
            _reqif_add_attr(vals, "AD-Status", entity.status)
            _reqif_add_attr(vals, "AD-ReqID", getattr(entity, "req_id", "") or "")
            _reqif_add_attr(vals, "AD-Body", _strip_html(getattr(entity, "body", None)))

        so_type = ET.SubElement(so, "TYPE")
        ET.SubElement(so_type, "SPEC-OBJECT-TYPE-REF").text = "SOT-Entity"

    # ── Specification with hierarchy ──────────────────────────────
    specifications = ET.SubElement(content, "SPECIFICATIONS")
    spec = ET.SubElement(
        specifications, "SPECIFICATION",
        IDENTIFIER=f"SPEC-{project.id}", LONG_NAME=project.name,
    )
    spec_st = ET.SubElement(spec, "TYPE")
    ET.SubElement(spec_st, "SPECIFICATION-TYPE-REF").text = "ST-Spec"

    children_el = ET.SubElement(spec, "CHILDREN")
    _build_reqif_hierarchy(children_el, project.id)

    # ── Write out ─────────────────────────────────────────────────
    et = ET.ElementTree(root)
    ET.indent(et, space="  ")
    et.write(filepath, encoding="utf-8", xml_declaration=True)


def _reqif_add_attr(parent, attr_def_id: str, value: str):
    """Add an ATTRIBUTE-VALUE-STRING element to a VALUES container."""
    if not value:
        return
    av = ET.SubElement(parent, "ATTRIBUTE-VALUE-STRING", THE_VALUE=value)
    defn = ET.SubElement(av, "DEFINITION")
    ET.SubElement(defn, "ATTRIBUTE-DEFINITION-STRING-REF").text = attr_def_id


def _build_reqif_hierarchy(parent_el, entity_id: int):
    """Recursively build SPEC-HIERARCHY elements for the ReqIF tree."""
    children = get_children(entity_id)
    for child in children:
        sh = ET.SubElement(
            parent_el, "SPEC-HIERARCHY",
            IDENTIFIER=f"SH-{child.id}",
        )
        obj_ref = ET.SubElement(sh, "OBJECT")
        ET.SubElement(obj_ref, "SPEC-OBJECT-REF").text = f"SO-{child.id}"

        grandchildren = get_children(child.id)
        if grandchildren:
            children_el = ET.SubElement(sh, "CHILDREN")
            for gc in grandchildren:
                _build_reqif_child(children_el, gc)


def _build_reqif_child(parent_el, entity):
    """Recursively build a single SPEC-HIERARCHY node and its children."""
    sh = ET.SubElement(
        parent_el, "SPEC-HIERARCHY",
        IDENTIFIER=f"SH-{entity.id}",
    )
    obj_ref = ET.SubElement(sh, "OBJECT")
    ET.SubElement(obj_ref, "SPEC-OBJECT-REF").text = f"SO-{entity.id}"

    grandchildren = get_children(entity.id)
    if grandchildren:
        children_el = ET.SubElement(sh, "CHILDREN")
        for gc in grandchildren:
            _build_reqif_child(children_el, gc)


# ═══════════════════════════════════════════════════════════════════
# DISPATCHER — called by the GUI
# ═══════════════════════════════════════════════════════════════════

# Maps format key → (file filter string, export function)
EXPORT_FORMATS = {
    "PDF":   ("PDF Files (*.pdf)",   export_pdf),
    "TXT":   ("Text Files (*.txt)",  export_txt),
    "DOCX":  ("Word Documents (*.docx)", export_docx),
    "CSV":   ("CSV Files (*.csv)",   export_csv),
    "ReqIF": ("ReqIF Files (*.reqif)", export_reqif),
}
